#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
A股盘前情报报告

功能：
1. 读取父目录 output 中最新的 V2.8.5_每日行情输出*.xlsx；
2. 抓取隔夜市场行情；
3. 抓取新闻源；
4. 按 V2.8.5 投资纪律生成盘前决策简报。

原则：
- 只输出“观察 / 暂停 / 复审 / 可进入买点检查”；
- 不输出直接买入建议；
- 不连接券商账户，不自动交易。
"""

from __future__ import annotations

import csv
import html
import json
import re
import sys
import xml.etree.ElementTree as ET
from dataclasses import dataclass
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any
from urllib.parse import quote, urljoin
from zoneinfo import ZoneInfo

import pandas as pd
try:
    import requests
except ImportError:  # 离线校验不依赖网络库
    requests = None

try:
    from dateutil import parser as date_parser
except ImportError:  # 离线校验不解析网络时间
    date_parser = None
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor

try:
    import feedparser
except ImportError:  # 允许先运行语法检查；正式使用请 pip install -r requirements.txt
    feedparser = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None


# ==================== 路径配置 ====================
PROJECT_DIR = Path(__file__).resolve().parent
FRAMEWORK_DIR = PROJECT_DIR.parent
DASHBOARD_PATTERNS = ["V2.8.5_每日行情输出*.xlsx", "V2.8.4_每日行情输出*.xlsx"]
FRAMEWORK_VERSION = "V2.8.5"
NEWS_SOURCES_CSV = PROJECT_DIR / "news_sources.csv"
OUTPUT_DIR = PROJECT_DIR / "output"
SHANGHAI_TZ = ZoneInfo("Asia/Shanghai")
POSITION_RISK_SHEET_NAMES = ("04_持仓风险", "05_持仓风险", "Positions_Action")
FIRST_YEAR_SHEET_NAMES = ("05_年度配置", "06_年度配置", "FirstYear_Allocation")
PORTFOLIO_SUMMARY_SHEET_NAMES = ("06_组合总览", "07_组合总览", "Dashboard")


# ==================== 隔夜行情配置 ====================
# 使用 Yahoo Finance 免费行情接口，失败时切换 Yahoo 备用 host，再切到 Stooq 历史 CSV。
# 任何入口失败都会写入审计；不会用旧数据或空数据补成“最新”。
YAHOO_CHART_HOSTS = ("query1.finance.yahoo.com", "query2.finance.yahoo.com")
MARKET_ITEMS = [
    {"name": "纳指", "symbols": ["^IXIC"], "stooq_symbols": ["^comp", "^ndq"], "unit": "点", "kind": "index"},
    {"name": "标普500", "symbols": ["^GSPC"], "stooq_symbols": ["^spx"], "unit": "点", "kind": "index"},
    {"name": "道指", "symbols": ["^DJI"], "stooq_symbols": ["^dji"], "unit": "点", "kind": "index"},
    {"name": "A50", "symbols": ["CN=F", "XIN9.FGI", "2823.HK"], "stooq_symbols": ["cn.f"], "unit": "点", "kind": "index"},
    {"name": "美债10Y", "symbols": ["^TNX"], "stooq_symbols": ["10usy.b"], "unit": "%", "kind": "yield"},
    {"name": "美元", "symbols": ["DX-Y.NYB", "DX=F"], "stooq_symbols": ["dx.f"], "unit": "点", "kind": "index"},
    {"name": "黄金", "symbols": ["GC=F", "XAUUSD=X"], "stooq_symbols": ["gc.f", "xauusd"], "unit": "美元/盎司", "kind": "commodity"},
    {"name": "原油", "symbols": ["CL=F", "BZ=F"], "stooq_symbols": ["cl.f", "brn.f"], "unit": "美元/桶", "kind": "commodity"},
]


# 原入口返回空结果或陈旧结果时按顺序尝试。备用入口仍然只做信息抓取，
# 失败会完整写入报告的数据审计章节，不会静默吞掉。
NEWS_SOURCE_FALLBACKS = {
    "MarketWatch": [
        {
            "Type": "rss",
            "URL": "https://news.google.com/rss/search?q=site%3Amarketwatch.com%20%28markets%20OR%20stocks%20OR%20China%29&hl=en-US&gl=US&ceid=US:en",
            "Label": "Google News site:marketwatch.com",
        }
    ],
    "CNBC": [
        {
            "Type": "rss",
            "URL": "https://news.google.com/rss/search?q=site%3Acnbc.com%20%28markets%20OR%20stocks%20OR%20China%20OR%20Federal%20Reserve%29&hl=en-US&gl=US&ceid=US:en",
            "Label": "Google News site:cnbc.com",
        }
    ],
    "财联社": [
        {
            "Type": "cls_api",
            "URL": "https://www.cls.cn/nodeapi/telegraphList?app=CailianpressWeb&category=&lastTime=&last_time=&os=web&refresh_type=1&rn=20&sv=7.7.5",
            "Label": "财联社电报API",
        },
        {
            "Type": "rss",
            "URL": "https://news.google.com/rss/search?q=site%3Acls.cn%20%28A%E8%82%A1%20OR%20%E7%BB%8F%E6%B5%8E%20OR%20%E6%94%BF%E7%AD%96%20OR%20%E5%B8%82%E5%9C%BA%29&hl=zh-CN&gl=CN&ceid=CN:zh-Hans",
            "Label": "Google News site:cls.cn",
        }
    ],
    "证券时报": [
        {
            "Type": "html",
            "URL": "https://www.stcn.com/kuaixun/",
            "Label": "证券时报快讯页",
        }
    ],
    "中国政府网": [
        {
            "Type": "html",
            "URL": "https://www.gov.cn/",
            "Label": "中国政府网首页要闻",
        }
    ],
    "央行": [
        {
            "Type": "html",
            "URL": "https://www.pbc.gov.cn/goutongjiaoliu/113456/113469/index.html",
            "Label": "央行 HTTPS 新闻页",
        }
    ],
    "证监会": [
        {
            "Type": "html",
            "URL": "https://www.csrc.gov.cn/csrc/index.shtml",
            "Label": "证监会 HTTPS 首页",
        }
    ],
    "发改委": [
        {
            "Type": "html",
            "URL": "https://www.ndrc.gov.cn/xwdt/",
            "Label": "发改委新闻动态首页",
        }
    ],
    "工信部": [
        {
            "Type": "html",
            "URL": "https://www.miit.gov.cn/xwdt/index.html",
            "Label": "工信部新闻动态首页",
        }
    ],
}


# ==================== 主题关键词 ====================
THEME_KEYWORDS = {
    "半导体": ["半导体", "芯片", "晶圆", "光刻", "封装", "先进制程", "semiconductor", "chip", "nvidia", "asml", "tsmc"],
    "AI算力": ["AI", "人工智能", "算力", "GPU", "数据中心", "云计算", "大模型", "openai", "nvidia", "data center"],
    "机器人": ["机器人", "具身智能", "人形机器人", "工业机器人", "robot", "humanoid"],
    "创新药": ["创新药", "医药", "生物医药", "药品", "医保", "FDA", "biotech", "pharma"],
    "新能源": ["新能源", "电池", "储能", "锂电", "光伏", "绿电", "电动车", "EV", "battery"],
    "宏观": ["美联储", "降息", "加息", "通胀", "美元", "美债", "财政", "货币政策", "关税", "Fed", "tariff", "inflation"],
}

NEGATIVE_WORDS = ["下跌", "暴跌", "制裁", "调查", "暂停", "收紧", "风险", "亏损", "降级", "禁令", "跌", "slump", "probe", "ban", "risk"]
POSITIVE_WORDS = ["上涨", "增长", "放松", "支持", "扩产", "突破", "批准", "创新高", "涨", "surge", "approve", "support", "record"]
ALLOWED_ACTIONS = {"观察", "暂停", "复审", "可进入买点检查"}
SKIP_TITLE_KEYWORDS = [
    "ICP备",
    "公网安备",
    "许可证",
    "版权所有",
    "沪金信备",
    "English Version",
    "网站地图",
    "客户端",
    "手机版",
    "登录",
    "注册",
]


@dataclass
class MarketQuote:
    name: str
    latest: float | None
    previous: float | None
    pct_change: float | None
    change: float | None
    unit: str
    symbol: str
    note: str = ""
    data_time: str = ""


@dataclass
class NewsItem:
    title: str
    time: str
    source: str
    group: str
    link: str
    summary: str


def clean_text(value: Any, max_len: int | None = None) -> str:
    """清理网页文本，去掉多余空格和 HTML 标签。"""
    if value is None:
        return ""
    text = html.unescape(str(value))
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    if max_len and len(text) > max_len:
        return text[: max_len - 1] + "…"
    return text


def to_float(value: Any) -> float | None:
    """安全转成数字；失败时返回 None。"""
    if value is None or pd.isna(value):
        return None
    try:
        return float(value)
    except (TypeError, ValueError):
        return None


def format_number(value: float | None, digits: int = 2) -> str:
    if value is None:
        return "未取到"
    return f"{value:,.{digits}f}"


def format_pct(value: float | None) -> str:
    if value is None:
        return "未取到"
    sign = "+" if value > 0 else ""
    return f"{sign}{value:.2f}%"


def safe_action(action: str) -> str:
    """所有动作都压到四个允许词内，避免生成直接交易建议。"""
    if action in ALLOWED_ACTIONS:
        return action
    if "复审" in action:
        return "复审"
    if "暂停" in action or "禁止" in action or "不" in action:
        return "暂停"
    if "检查" in action:
        return "可进入买点检查"
    return "观察"


# ==================== 读取 V2.8.5 每日行情输出 ====================
def find_latest_dashboard_xlsx() -> Path | None:
    """查找 output 文件夹中最新的每日行情 Excel。"""
    dashboard_dir = FRAMEWORK_DIR / "output"
    candidates: list[Path] = []
    for pattern in DASHBOARD_PATTERNS:
        matches = [path for path in dashboard_dir.glob(pattern) if not path.name.startswith("~$")]
        if matches:
            candidates = matches
            break
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def get_sheet_with_fallback(
    sheets: dict[str, pd.DataFrame],
    *names: str,
) -> pd.DataFrame:
    """优先返回新行动页；旧版工作簿名只用于兼容回退。"""
    for name in names:
        frame = sheets.get(name)
        if frame is not None and not frame.empty:
            return frame
    for name in names:
        if name in sheets:
            return sheets[name]
    return pd.DataFrame()


def read_dashboard_workbook() -> dict[str, pd.DataFrame]:
    """读取每日行情输出里的关键工作表。"""
    dashboard_xlsx = find_latest_dashboard_xlsx()
    if dashboard_xlsx is None:
        print(f"【提示】找不到每日行情输出：{FRAMEWORK_DIR / 'output'}")
        return {}

    print(f"已读取每日行情文件：{dashboard_xlsx}")
    wanted_sheets = [
        "01_今日决策", "02_今日动作", "03_买入候选",
        "04_持仓风险", "05_持仓风险", "05_年度配置", "06_年度配置", "06_组合总览", "07_组合总览",
        "Decision_Center", "Dashboard", "Market_Data", "Double_Anchor", "Emotion",
        "Quality_Score", "Exposure", "Buy_Filter", "Positions", "Positions_Action",
        "FirstYear_Allocation", "Execution_Plan", "Checks", "Framework_Rules",
        "长期跟踪个股", "Investment_Profile", "Data_Sources",
    ]
    sheets: dict[str, pd.DataFrame] = {}
    excel_file = pd.ExcelFile(dashboard_xlsx)
    for sheet_name in wanted_sheets:
        if sheet_name in excel_file.sheet_names:
            sheets[sheet_name] = pd.read_excel(dashboard_xlsx, sheet_name=sheet_name, dtype={"Code": str})
        else:
            sheets[sheet_name] = pd.DataFrame()
    return sheets


# ==================== 抓取隔夜行情 ====================
def yahoo_chart_url(symbol: str, host: str = YAHOO_CHART_HOSTS[0]) -> str:
    encoded = quote(symbol, safe="")
    return f"https://{host}/v8/finance/chart/{encoded}?range=7d&interval=1d"


def stooq_history_url(symbol: str) -> str:
    encoded = quote(symbol, safe=".^")
    end = datetime.now(SHANGHAI_TZ).strftime("%Y%m%d")
    start = (datetime.now(SHANGHAI_TZ) - timedelta(days=21)).strftime("%Y%m%d")
    return f"https://stooq.com/q/d/l/?s={encoded}&d1={start}&d2={end}&i=d"


def timestamp_to_shanghai(value: Any) -> str:
    """Convert an epoch timestamp to an auditable Shanghai-time label."""
    try:
        return datetime.fromtimestamp(float(value), SHANGHAI_TZ).strftime("%Y-%m-%d %H:%M")
    except (TypeError, ValueError, OSError):
        return ""


def fetch_one_yahoo_symbol(
    session: requests.Session,
    item: dict[str, Any],
    symbol: str,
    host: str = YAHOO_CHART_HOSTS[0],
) -> MarketQuote | None:
    """抓取一个 Yahoo symbol 的最近两个交易日收盘数据。"""
    response = session.get(yahoo_chart_url(symbol, host), timeout=12)
    response.raise_for_status()
    data = response.json()
    result = data.get("chart", {}).get("result", [])
    if not result:
        return None

    result0 = result[0]
    meta = result0.get("meta", {})
    timestamps = result0.get("timestamp", [])
    closes = result0.get("indicators", {}).get("quote", [{}])[0].get("close", [])
    valid_closes = [float(x) for x in closes if x is not None]
    if not valid_closes:
        return None

    latest_raw = to_float(meta.get("regularMarketPrice")) or valid_closes[-1]
    previous_raw = to_float(meta.get("chartPreviousClose"))
    if previous_raw is None and len(valid_closes) >= 2:
        previous_raw = valid_closes[-2]

    # Yahoo 的 ^TNX 有时返回 4.49，有时历史接口会返回 44.9；
    # 这里动态判断，避免把 4.49% 错写成 0.449%。
    divisor = float(item.get("divisor", 1.0))
    if item.get("kind") == "yield":
        raw_for_check = latest_raw if latest_raw is not None else previous_raw
        if raw_for_check is not None and raw_for_check > 20:
            divisor = 10.0
    latest = latest_raw / divisor
    previous = previous_raw / divisor if previous_raw is not None else None
    change = latest - previous if previous is not None else None

    if previous is not None and previous != 0:
        pct_change = change / previous * 100
    else:
        pct_change = None

    return MarketQuote(
        name=item["name"],
        latest=round(latest, 4),
        previous=round(previous, 4) if previous is not None else None,
        pct_change=round(pct_change, 4) if pct_change is not None else None,
        change=round(change, 4) if change is not None else None,
        unit=item["unit"],
        symbol=f"Yahoo:{symbol}",
        data_time=timestamp_to_shanghai(meta.get("regularMarketTime") or (timestamps[-1] if timestamps else None)),
    )


def fetch_one_stooq_symbol(session: requests.Session, item: dict[str, Any], symbol: str) -> MarketQuote | None:
    """抓取 Stooq 日线 CSV；只用最近两个有效收盘价，不估算缺口。"""
    response = session.get(stooq_history_url(symbol), timeout=12)
    response.raise_for_status()
    rows = list(csv.DictReader(response.text.splitlines()))
    valid_rows = [
        row for row in rows
        if to_float(row.get("Close")) is not None and clean_text(row.get("Date"))
    ]
    if not valid_rows:
        return None

    latest_row = valid_rows[-1]
    previous_row = valid_rows[-2] if len(valid_rows) >= 2 else None
    latest_raw = to_float(latest_row.get("Close"))
    previous_raw = to_float(previous_row.get("Close")) if previous_row else None
    if latest_raw is None:
        return None

    divisor = float(item.get("divisor", 1.0))
    if item.get("kind") == "yield":
        raw_for_check = latest_raw if latest_raw is not None else previous_raw
        if raw_for_check is not None and raw_for_check > 20:
            divisor = 10.0
    latest = latest_raw / divisor
    previous = previous_raw / divisor if previous_raw is not None else None
    change = latest - previous if previous is not None else None
    pct_change = change / previous * 100 if previous not in (None, 0) else None

    return MarketQuote(
        name=item["name"],
        latest=round(latest, 4),
        previous=round(previous, 4) if previous is not None else None,
        pct_change=round(pct_change, 4) if pct_change is not None else None,
        change=round(change, 4) if change is not None else None,
        unit=item["unit"],
        symbol=f"Stooq:{symbol}",
        data_time=clean_text(latest_row.get("Date")),
    )


def market_attempts_for_item(item: dict[str, Any]) -> list[dict[str, Any]]:
    attempts: list[dict[str, Any]] = []
    for symbol in item.get("symbols", []):
        for host in YAHOO_CHART_HOSTS:
            attempts.append({"provider": "yahoo", "host": host, "symbol": symbol, "label": f"Yahoo {host} {symbol}"})
    for symbol in item.get("stooq_symbols", []):
        attempts.append({"provider": "stooq", "symbol": symbol, "label": f"Stooq {symbol}"})
    return attempts


def fetch_market_attempt(session: requests.Session, item: dict[str, Any], attempt: dict[str, Any]) -> MarketQuote | None:
    if attempt["provider"] == "yahoo":
        return fetch_one_yahoo_symbol(session, item, attempt["symbol"], attempt["host"])
    if attempt["provider"] == "stooq":
        return fetch_one_stooq_symbol(session, item, attempt["symbol"])
    raise ValueError(f"Unknown market provider: {attempt['provider']}")


def format_market_attempt_summary(attempts: list[tuple[str, str]], max_attempts: int = 5) -> str:
    visible = attempts[:max_attempts]
    details = "；".join(f"{label}：{result}" for label, result in visible)
    if len(attempts) > max_attempts:
        details += f"；另{len(attempts) - max_attempts}次尝试失败"
    return details


def fetch_market_quotes() -> list[MarketQuote]:
    """抓取隔夜行情；每个指标成功一个 symbol 即停止。"""
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": "Mozilla/5.0 premarket-report/1.0",
            "Accept": "application/json,text/plain,*/*",
        }
    )

    quotes: list[MarketQuote] = []
    for item in MARKET_ITEMS:
        quote_obj = None
        attempts: list[tuple[str, str]] = []
        for attempt in market_attempts_for_item(item):
            try:
                quote_obj = fetch_market_attempt(session, item, attempt)
                if quote_obj:
                    if attempts:
                        attempts.append((attempt["label"], "备用源成功"))
                        quote_obj.note = "行情源追踪：" + format_market_attempt_summary(attempts)
                    break
            except Exception as exc:
                attempts.append((attempt["label"], f"失败 {clean_text(exc, 80)}"))
                continue
            attempts.append((attempt["label"], "无有效收盘数据"))
        if quote_obj is None:
            quote_obj = MarketQuote(
                name=item["name"],
                latest=None,
                previous=None,
                pct_change=None,
                change=None,
                unit=item["unit"],
                symbol="/".join(item.get("symbols", []) + item.get("stooq_symbols", [])),
                note=f"抓取失败：{format_market_attempt_summary(attempts)}",
            )
        quotes.append(quote_obj)
    return quotes


def judge_risk_preference(quotes: list[MarketQuote]) -> tuple[str, list[str]]:
    """用隔夜权益、A50、美元、美债、黄金、原油做第一版风险偏好判断。"""
    by_name = {q.name: q for q in quotes}
    score = 0
    notes: list[str] = []

    us_pcts = [by_name[name].pct_change for name in ["纳指", "标普500", "道指"] if by_name.get(name) and by_name[name].pct_change is not None]
    if us_pcts:
        us_avg = sum(us_pcts) / len(us_pcts)
        if us_avg >= 0.3:
            score += 1
            notes.append("美股三大指数整体偏强")
        elif us_avg <= -0.3:
            score -= 1
            notes.append("美股三大指数整体偏弱")

    a50 = by_name.get("A50")
    if a50 and a50.pct_change is not None:
        if a50.pct_change >= 0.3:
            score += 1
            notes.append("A50 对 A股开盘情绪有正向牵引")
        elif a50.pct_change <= -0.3:
            score -= 1
            notes.append("A50 对 A股开盘情绪偏负面")

    usd = by_name.get("美元")
    if usd and usd.pct_change is not None:
        if usd.pct_change >= 0.3:
            score -= 1
            notes.append("美元走强压制风险资产")
        elif usd.pct_change <= -0.3:
            score += 1
            notes.append("美元走弱有利于风险偏好")

    treasury = by_name.get("美债10Y")
    if treasury and treasury.change is not None:
        bps_change = treasury.change * 100
        if bps_change >= 3:
            score -= 1
            notes.append("美债10Y上行，对成长资产不利")
        elif bps_change <= -3:
            score += 1
            notes.append("美债10Y回落，对成长资产压力下降")

    gold = by_name.get("黄金")
    if gold and gold.pct_change is not None and gold.pct_change >= 0.8 and score <= 0:
        score -= 1
        notes.append("黄金明显走强，可能反映避险升温")

    oil = by_name.get("原油")
    if oil and oil.pct_change is not None and oil.pct_change >= 1.5:
        score -= 1
        notes.append("原油明显上涨，通胀预期压力需观察")

    if score >= 2:
        return "上升", notes
    if score <= -1:
        return "下降", notes
    return "中性", notes or ["隔夜变量没有形成明显单边信号"]


# ==================== 抓取新闻 ====================
def load_news_sources() -> list[dict[str, Any]]:
    """读取 news_sources.csv。"""
    if not NEWS_SOURCES_CSV.exists():
        return []
    sources: list[dict[str, Any]] = []
    with NEWS_SOURCES_CSV.open(encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row["Limit"] = int(row.get("Limit") or 5)
            sources.append(row)
    return sources


def parse_time(value: Any) -> str:
    """把新闻时间转成字符串；失败时返回空。"""
    if not value:
        return ""
    try:
        normalized = clean_text(value, 80).replace("年", "-").replace("月", "-").replace("日", "")
        dt = date_parser.parse(normalized)
        if dt.tzinfo is not None:
            dt = dt.astimezone(SHANGHAI_TZ)
        return dt.strftime("%Y-%m-%d %H:%M")
    except Exception:
        return clean_text(value, 30)


def parse_datetime_or_none(value: Any) -> datetime | None:
    """解析新闻日期；支持常见中文日期，失败返回 None。"""
    if not value:
        return None
    text = clean_text(value, 40)
    text = text.replace("年", "-").replace("月", "-").replace("日", "")
    try:
        dt = date_parser.parse(text)
        if dt.tzinfo is not None:
            dt = dt.astimezone(SHANGHAI_TZ).replace(tzinfo=None)
        return dt
    except Exception:
        return None


def is_recent_news(item: NewsItem, max_days: int = 3) -> bool:
    """只过滤明确可解析且过旧的新闻；无法解析时间的网页新闻先保留。"""
    dt = parse_datetime_or_none(item.time)
    if dt is None:
        return True
    return dt >= datetime.now() - timedelta(days=max_days)


def infer_news_time_from_url(url: str) -> str:
    """Infer only an explicit YYYYMMDD date embedded in a source URL."""
    match = re.search(r"(?<!\d)(20\d{6})(?:\d{0,12})?(?!\d)", url or "")
    if match is None:
        match = re.search(r"t(20\d{6})(?:_|\D)", url or "")
    if match is None:
        return ""
    try:
        return datetime.strptime(match.group(1), "%Y%m%d").strftime("%Y-%m-%d")
    except ValueError:
        return ""


def extract_article_time(session: requests.Session, url: str) -> str:
    """Read publication time from article metadata/text without inventing it."""
    if not url or not url.lower().startswith(("http://", "https://")):
        return ""
    response = session.get(url, timeout=10)
    response.raise_for_status()
    soup = BeautifulSoup(response.content, "html.parser")
    meta_keys = {
        "article:published_time", "pubdate", "publishdate", "publish_time",
        "date", "dc.date", "dcterms.date", "sailthru.date",
    }
    for tag in soup.find_all("meta"):
        key = clean_text(tag.get("property") or tag.get("name") or "").lower()
        content = clean_text(tag.get("content") or "", 80)
        if key in meta_keys and content:
            parsed = parse_time(content)
            if parsed:
                return parsed
    text = clean_text(soup.get_text(" "), 5000)
    return parse_time(find_date_in_text(text)) if text else ""


def enrich_missing_news_times(session: requests.Session, items: list[NewsItem]) -> list[NewsItem]:
    """Fill missing times from explicit URL dates, then source article metadata."""
    for item in items:
        if item.time:
            continue
        item.time = infer_news_time_from_url(item.link)
        if item.time:
            continue
        try:
            item.time = extract_article_time(session, item.link)
        except Exception:
            item.time = ""
    return items


def fetch_rss_source(session: requests.Session, source: dict[str, Any]) -> list[NewsItem]:
    """抓取 RSS 新闻源。"""
    response = session.get(source["URL"], timeout=15)
    response.raise_for_status()
    items: list[NewsItem] = []

    if feedparser is not None:
        parsed = feedparser.parse(response.content)
        for entry in parsed.entries[: source["Limit"]]:
            title = clean_text(entry.get("title", ""), 120)
            link = entry.get("link", "")
            if not title or should_skip_title(title, link):
                continue
            summary = clean_text(entry.get("summary", "") or entry.get("description", ""), 220)
            published = entry.get("published") or entry.get("updated") or ""
            items.append(
                NewsItem(
                    title=title,
                    time=parse_time(published),
                    source=source["Source"],
                    group=source["Group"],
                    link=link,
                    summary=summary or title,
                )
            )
        return items

    # 备用解析：不依赖 feedparser，只解析常见 RSS/Atom 字段。
    root = ET.fromstring(response.content)
    rss_entries = root.findall(".//item")
    atom_entries = root.findall(".//{http://www.w3.org/2005/Atom}entry")

    for entry in (rss_entries or atom_entries)[: source["Limit"]]:
        if rss_entries:
            title = clean_text(entry.findtext("title"), 120)
            link = clean_text(entry.findtext("link"))
            summary = clean_text(entry.findtext("description"), 220)
            published = entry.findtext("pubDate")
        else:
            title = clean_text(entry.findtext("{http://www.w3.org/2005/Atom}title"), 120)
            link_tag = entry.find("{http://www.w3.org/2005/Atom}link")
            link = clean_text(link_tag.attrib.get("href") if link_tag is not None else "")
            summary = clean_text(entry.findtext("{http://www.w3.org/2005/Atom}summary"), 220)
            published = entry.findtext("{http://www.w3.org/2005/Atom}updated")

        if not title or should_skip_title(title, link):
            continue
        items.append(
            NewsItem(
                title=title,
                time=parse_time(published),
                source=source["Source"],
                group=source["Group"],
                link=link,
                summary=summary or title,
            )
        )
    return items


def find_date_in_text(text: str) -> str:
    """从网页文本中尝试找日期。"""
    match = re.search(r"(20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}日?\s*\d{0,2}:?\d{0,2})", text)
    return clean_text(match.group(1), 30) if match else ""


def fetch_html_source(session: requests.Session, source: dict[str, Any]) -> list[NewsItem]:
    """通用 HTML 新闻列表抓取。适合官方源和普通新闻首页。"""
    if BeautifulSoup is None:
        raise RuntimeError("缺少 beautifulsoup4，请先运行 pip install -r requirements.txt")

    response = session.get(source["URL"], timeout=15)
    response.raise_for_status()
    # Let BeautifulSoup honor the page's own charset declaration. requests'
    # apparent_encoding misclassifies some Chinese government pages.
    soup = BeautifulSoup(response.content, "html.parser")

    items: list[NewsItem] = []
    seen_titles: set[str] = set()
    for link_tag in soup.find_all("a"):
        title = clean_text(link_tag.get_text(" "), 120)
        href = link_tag.get("href")
        if not title or not href:
            continue
        full_link = urljoin(source["URL"], href)
        if should_skip_title(title, full_link) or title in seen_titles:
            continue

        parent_text = clean_text(link_tag.parent.get_text(" ") if link_tag.parent else "", 220)
        news_time = parse_time(find_date_in_text(parent_text))
        summary = parent_text if parent_text and parent_text != title else title
        seen_titles.add(title)
        items.append(
            NewsItem(
                title=title,
                time=news_time,
                source=source["Source"],
                group=source["Group"],
                link=full_link,
                summary=summary,
            )
        )
        if len(items) >= source["Limit"]:
            break
    return items


def fetch_cls_api(session: requests.Session, source: dict[str, Any]) -> list[NewsItem]:
    """财联社电报接口。接口偶尔会变化，失败时会被主流程记录。"""
    response = session.get(source["URL"], timeout=15)
    if response.status_code >= 400:
        fallback_source = source.copy()
        fallback_source["URL"] = "https://www.cls.cn/telegraph"
        return fetch_html_source(session, fallback_source)
    response.raise_for_status()
    data = response.json()
    candidates = data.get("data", {}).get("roll_data") or data.get("data", {}).get("telegraphs") or []

    items: list[NewsItem] = []
    for item in candidates[: source["Limit"]]:
        title = clean_text(item.get("title") or item.get("content") or "", 120)
        link = item.get("url") or item.get("shareurl") or source["URL"]
        if not title or should_skip_title(title, link):
            continue
        ctime = item.get("ctime") or item.get("time") or item.get("modified_time")
        news_time = ""
        if isinstance(ctime, (int, float)):
            news_time = datetime.fromtimestamp(ctime).strftime("%Y-%m-%d %H:%M")
        elif ctime:
            news_time = parse_time(ctime)
        items.append(
            NewsItem(
                title=title,
                time=news_time,
                source=source["Source"],
                group=source["Group"],
                link=link,
                summary=clean_text(item.get("brief") or item.get("content") or title, 220),
            )
        )
    return items


def fetch_news_source_once(session: requests.Session, source: dict[str, Any]) -> list[NewsItem]:
    source_type = (source.get("Type") or "").strip().lower()
    if source_type == "rss":
        return fetch_rss_source(session, source)
    if source_type == "cls_api":
        return fetch_cls_api(session, source)
    return fetch_html_source(session, source)


def format_source_attempt_summary(source_name: str, attempts: list[tuple[str, str]]) -> str:
    details = "；".join(f"{label}：{result}" for label, result in attempts)
    return f"{source_name}源审计：{details}"


def fetch_news() -> tuple[list[NewsItem], list[str]]:
    """抓取所有新闻源；单个来源失败不影响报告生成。"""
    session = requests.Session()
    session.headers.update(
        {
            "User-Agent": "Mozilla/5.0 premarket-report/1.0",
            "Accept": "text/html,application/xhtml+xml,application/xml,application/rss+xml,application/json,*/*",
        }
    )
    news: list[NewsItem] = []
    warnings: list[str] = []

    for source in load_news_sources():
        source_name = source.get("Source", "未知来源")
        attempts: list[tuple[str, str]] = []
        selected_items: list[NewsItem] = []

        # Original source gets one retry. A successful HTTP response with only
        # stale/invalid rows is also considered a failed attempt.
        for attempt_no in range(1, 3):
            try:
                items = fetch_news_source_once(session, source)
                items = enrich_missing_news_times(session, items)
                items = [item for item in items if is_recent_news(item)]
                if items:
                    selected_items = items
                    if attempt_no > 1:
                        attempts.append((f"原入口重试{attempt_no}", f"成功，{len(items)}条"))
                    break
                attempts.append((f"原入口尝试{attempt_no}", "HTTP/解析成功但0条近3日有效新闻"))
            except Exception as exc:
                attempts.append((f"原入口尝试{attempt_no}", f"失败 {clean_text(exc, 100)}"))

        if not selected_items:
            for fallback in NEWS_SOURCE_FALLBACKS.get(source_name, []):
                fallback_source = source.copy()
                fallback_source.update(fallback)
                label = fallback.get("Label") or fallback.get("URL", "备用源")
                try:
                    items = fetch_news_source_once(session, fallback_source)
                    items = enrich_missing_news_times(session, items)
                    items = [item for item in items if is_recent_news(item)]
                    if items:
                        selected_items = items
                        attempts.append((label, f"备用源成功，{len(items)}条"))
                        break
                    attempts.append((label, "备用源可访问但0条近3日有效新闻"))
                except Exception as exc:
                    attempts.append((label, f"备用源失败 {clean_text(exc, 100)}"))

        if attempts:
            warnings.append(format_source_attempt_summary(source_name, attempts))
        if not selected_items and not attempts:
            warnings.append(f"{source_name}源审计：未配置可用尝试")
        news.extend(selected_items)

    # 去重：优先按链接，其次按标题。
    deduped: list[NewsItem] = []
    seen: set[str] = set()
    for item in news:
        key = item.link or item.title
        if key in seen:
            continue
        seen.add(key)
        deduped.append(item)

    missing_time_by_source: dict[str, int] = {}
    for item in deduped:
        if not item.time:
            missing_time_by_source[item.source] = missing_time_by_source.get(item.source, 0) + 1
    for source_name, count in missing_time_by_source.items():
        warnings.append(f"{source_name}时间字段：{count}条经URL日期与原文元数据查询后仍未取得")

    return deduped, warnings


# ==================== 框架分析 ====================
def text_contains_any(text: str, keywords: list[str]) -> bool:
    lower = text.lower()
    return any(keyword.lower() in lower for keyword in keywords)


def should_skip_title(title: str, link: str = "") -> bool:
    """过滤网页上的备案、导航、版权等非新闻链接。"""
    if len(title) < 8:
        return True
    if any(keyword.lower() in title.lower() for keyword in SKIP_TITLE_KEYWORDS):
        return True
    if "beian" in link.lower() or "ifcert" in link.lower():
        return True
    return False


def classify_news_themes(item: NewsItem) -> list[str]:
    """给新闻打主题标签。"""
    text = f"{item.title} {item.summary}"
    return [theme for theme, keywords in THEME_KEYWORDS.items() if text_contains_any(text, keywords)]


def sentiment_from_news(items: list[NewsItem]) -> str:
    """简单判断相关新闻偏正面/负面/中性。"""
    text = " ".join([f"{item.title} {item.summary}" for item in items])
    neg = sum(1 for word in NEGATIVE_WORDS if word.lower() in text.lower())
    pos = sum(1 for word in POSITIVE_WORDS if word.lower() in text.lower())
    if neg > pos:
        return "偏负面"
    if pos > neg:
        return "偏正面"
    return "中性"


def get_anchor_gate(sheets: dict[str, pd.DataFrame], risk_preference: str) -> tuple[str, str]:
    """读取双锚综合灯号，并转成盘前动作门槛。"""
    df = sheets.get("Double_Anchor", pd.DataFrame())
    if df.empty or "锚点" not in df.columns:
        return "未读取到双锚", "观察"

    combined = df[df["锚点"].astype(str) == "综合灯号"]
    if combined.empty:
        return "未读取到综合灯号", "观察"

    light = str(combined.iloc[0].get("灯号", "观察"))
    if "红" in light or risk_preference == "下降":
        return light, "暂停"
    if light in {"双绿", "黄绿"} and risk_preference in {"上升", "中性"}:
        return light, "可进入买点检查"
    return light, "观察"


def get_buy_filter_risks(sheets: dict[str, pd.DataFrame]) -> list[str]:
    """读取买点过滤器中的风险项。"""
    df = get_sheet_with_fallback(sheets, "03_买入候选", "Buy_Filter")
    if df.empty:
        return ["未读取到 03_买入候选，盘中不得跳过买点过滤器"]

    risks: list[str] = []
    for _, row in df.iterrows():
        code = str(row.get("Code", "")).strip()
        name = str(row.get("Name", "")).strip()
        if code.upper().endswith((".SH", ".SZ", ".CSI")):
            continue
        veto = str(row.get("一票否决", "")).strip()
        pct = to_float(row.get("PctChg"))
        intraday_pos = to_float(row.get("日内位置"))
        if veto == "是":
            risks.append(f"{name}({code}) 一票否决：{clean_text(row.get('否决原因', '')) or '原因待确认'}")
        elif pct is not None and pct >= 3:
            risks.append(f"{name}({code}) 涨幅偏高：盘中禁止追高")
        elif intraday_pos is not None and intraday_pos >= 0.8:
            risks.append(f"{name}({code}) 日内位置偏高：需要等待回落验证")
    return risks[:8] or ["未发现一票否决项；盘中仍需逐项检查"]


def position_theme(name: str, role: str) -> str:
    """根据持仓名称和角色粗略归类主题。"""
    text = f"{name} {role}"
    if text_contains_any(text, THEME_KEYWORDS["机器人"]):
        return "机器人"
    if text_contains_any(text, THEME_KEYWORDS["AI算力"]):
        return "AI算力"
    if text_contains_any(text, THEME_KEYWORDS["半导体"]):
        return "半导体"
    if text_contains_any(text, THEME_KEYWORDS["创新药"]):
        return "创新药"
    if text_contains_any(text, THEME_KEYWORDS["新能源"]):
        return "新能源"
    return "宏观"


def build_position_impacts(
    sheets: dict[str, pd.DataFrame],
    news: list[NewsItem],
    risk_preference: str,
    anchor_action: str,
) -> list[dict[str, str]]:
    """生成持仓影响表。"""
    positions = sheets.get("Positions", pd.DataFrame())
    actions = get_sheet_with_fallback(sheets, *POSITION_RISK_SHEET_NAMES)
    if positions.empty:
        return []

    action_by_code: dict[str, pd.Series] = {}
    if not actions.empty and "Code" in actions.columns:
        action_by_code = {str(row["Code"]).strip(): row for _, row in actions.iterrows()}

    rows: list[dict[str, str]] = []
    for _, pos in positions.iterrows():
        code = str(pos.get("Code", "")).strip()
        name = str(pos.get("Name", "")).strip()
        role = str(pos.get("Role", "")).strip()
        theme = position_theme(name, role)
        theme_news = [item for item in news if theme in classify_news_themes(item)]
        related_titles = "；".join(item.title for item in theme_news[:2]) or "暂无明确相关新闻"

        framework = "按原纪律观察"
        action = "观察"

        weight = to_float(pos.get("Weight"))
        target_weight = to_float(pos.get("Target Weight"))
        annual_item = clean_text(pos.get("年度配置项", ""))
        annual_gap = to_float(pos.get("年度资金缺口"))
        annual_completion = to_float(pos.get("年度完成率"))
        if weight is not None and target_weight is not None and weight > target_weight:
            framework = f"证券账户仓位{weight:.2%}，超过目标{target_weight:.2%}"
            if annual_item and annual_gap is not None:
                framework += f"；第一年“{annual_item}”全资产缺口{annual_gap:,.0f}元，但不得据此突破账户上限"
            action = "复审"
        elif annual_item and annual_completion is not None:
            framework = f"第一年“{annual_item}”完成{annual_completion:.1%}；年度缺口仅作长期规划"

        action_row = action_by_code.get(code)
        if action_row is not None:
            raw_action = str(action_row.get("动作建议", ""))
            triggered = str(action_row.get("触发提醒", "")) == "是"
            if triggered or "复审" in raw_action:
                framework = clean_text(action_row.get("说明", "")) or "触发持仓纪律提醒"
                action = "复审"

        if action != "复审":
            if risk_preference == "下降" or anchor_action == "暂停":
                framework = "风险偏好或双锚不支持新动作"
                action = "暂停"
            elif theme_news:
                news_sentiment = sentiment_from_news(theme_news)
                framework = f"{theme}相关新闻情绪：{news_sentiment}"
                action = "复审" if news_sentiment == "偏负面" else "观察"

        rows.append(
            {
                "持仓": f"{name}({code})",
                "新闻/行情影响": clean_text(related_titles, 160),
                "框架判断": framework,
                "动作": safe_action(action),
            }
        )
    return rows


def build_theme_view(news: list[NewsItem], risk_preference: str, anchor_action: str) -> dict[str, str]:
    """生成四大主线判断。"""
    output: dict[str, str] = {}
    for theme in ["半导体", "AI算力", "机器人", "创新药"]:
        theme_news = [item for item in news if theme in classify_news_themes(item)]
        sentiment = sentiment_from_news(theme_news) if theme_news else "中性"

        if risk_preference == "下降" or anchor_action == "暂停":
            action = "暂停"
        elif sentiment == "偏负面":
            action = "复审"
        elif risk_preference == "上升" and anchor_action == "可进入买点检查":
            action = "可进入买点检查"
        else:
            action = "观察"

        title_part = f"相关新闻 {len(theme_news)} 条，情绪{sentiment}"
        output[theme] = f"{title_part}；动作：{safe_action(action)}"
    return output


def build_today_discipline(
    position_impacts: list[dict[str, str]],
    buy_filter_risks: list[str],
    risk_preference: str,
    anchor_action: str,
    theme_view: dict[str, str],
) -> dict[str, str]:
    """生成今日纪律。"""
    review_items = [row["持仓"] for row in position_impacts if row["动作"] == "复审"]
    observable_themes = [theme for theme, text in theme_view.items() if "观察" in text or "可进入买点检查" in text]

    if risk_preference == "下降" or anchor_action == "暂停":
        allowed = "成长类新动作暂停；仅保留观察和复审"
    elif observable_themes:
        allowed = "、".join(observable_themes)
    else:
        allowed = "仅观察"

    return {
        "禁止追高": "是；出现高开高走或日内位置偏高时，一律等待买点过滤器",
        "禁止自动挂单": "是；本项目不连接券商、不自动交易",
        "需要复审": "、".join(review_items) if review_items else "暂无持仓触发专项复审",
        "可观察": allowed,
        "买点过滤器风险": "；".join(buy_filter_risks[:5]),
    }


def build_framework_snapshot(sheets: dict[str, pd.DataFrame]) -> dict[str, str]:
    """Condense the V2.8.5 workbook into decision-critical premarket facts."""
    snapshot = {
        "双锚": "待确认",
        "情绪": "待确认",
        "情绪可靠性": "待确认",
        "标准首批": "0",
        "半额复核": "0",
        "质量待补": "待确认",
        "全资产权益仓位": "待确认",
        "第一年目标占比": "待确认",
        "第一年当前占比": "待确认",
        "第一年完成率": "待确认",
        "第一年资金缺口": "待确认",
        "现金安全垫": "待确认",
        "校验状态": "待确认",
        "数据状态": "条件式",
    }
    anchor = sheets.get("Double_Anchor", pd.DataFrame())
    if not anchor.empty and "锚点" in anchor.columns:
        row = anchor[anchor["锚点"].astype(str) == "综合灯号"]
        if not row.empty:
            snapshot["双锚"] = str(row.iloc[0].get("灯号", "待确认"))

    emotion = sheets.get("Emotion", pd.DataFrame())
    if not emotion.empty and "指标" in emotion.columns:
        row = emotion[emotion["指标"].astype(str) == "情绪温度"]
        if not row.empty:
            value = to_float(row.iloc[0].get("数值"))
            status = str(row.iloc[0].get("状态", ""))
            snapshot["情绪"] = f"{int(value)}级·{status}" if value is not None else status or "待确认"
            snapshot["情绪可靠性"] = str(row.iloc[0].get("可靠性", "待确认"))

    quality = sheets.get("Quality_Score", pd.DataFrame())
    if not quality.empty and "质量状态" in quality.columns:
        snapshot["质量待补"] = str(int((quality["质量状态"].astype(str) == "数据不足，需人工评分").sum()))

    buy_filter = get_sheet_with_fallback(sheets, "03_买入候选", "Buy_Filter")
    if not buy_filter.empty and "买点灯号" in buy_filter.columns:
        snapshot["标准首批"] = str(int((buy_filter["买点灯号"] == "绿").sum()))
        snapshot["半额复核"] = str(int((buy_filter["买点灯号"] == "黄").sum()))

    dashboard = get_sheet_with_fallback(sheets, *PORTFOLIO_SUMMARY_SHEET_NAMES)
    if not dashboard.empty and {"项目", "内容"}.issubset(dashboard.columns):
        dashboard_map = dict(zip(dashboard["项目"].astype(str), dashboard["内容"]))
        value = to_float(dashboard_map.get("全资产权益穿透仓位"))
        if value is not None:
            snapshot["全资产权益仓位"] = f"{value:.2%}"
        cash_status = dashboard_map.get("现金安全垫状态")
        cash_value = to_float(dashboard_map.get("全资产流动性安全垫代理"))
        if cash_status or cash_value is not None:
            snapshot["现金安全垫"] = f"{cash_status or '待确认'} · {cash_value:,.0f}元" if cash_value is not None else str(cash_status)
        first_target = to_float(dashboard_map.get("第一年配置目标占比"))
        first_current = to_float(dashboard_map.get("第一年当前已配置占比"))
        first_completion = to_float(dashboard_map.get("第一年配置完成率"))
        first_gap = to_float(dashboard_map.get("第一年资金缺口"))
        snapshot["第一年目标占比"] = f"{first_target:.1%}" if first_target is not None else "待确认"
        snapshot["第一年当前占比"] = f"{first_current:.1%}" if first_current is not None else "待确认"
        snapshot["第一年完成率"] = f"{first_completion:.1%}" if first_completion is not None else "待确认"
        snapshot["第一年资金缺口"] = f"{first_gap:,.0f}元" if first_gap is not None else "待确认"
        missing = to_float(dashboard_map.get("未取到行情数量"))
        snapshot["数据状态"] = "完整" if missing == 0 else f"{int(missing or 0)}项行情缺失，条件式"

    checks = sheets.get("Checks", pd.DataFrame())
    if not checks.empty and "状态" in checks.columns:
        non_ok = checks[~checks["状态"].astype(str).isin(["OK"])]
        snapshot["校验状态"] = "全部通过" if non_ok.empty else f"{len(non_ok)}项待处理"
    if snapshot.get("质量待补") not in {"0", "待确认"}:
        snapshot["数据状态"] = f"行情完整；{snapshot['质量待补']}只标的质量评分待补，决策条件式"
    return snapshot


# ==================== 生成 Word 报告 ====================
def news_lines(news: list[NewsItem], group: str, limit: int = 50) -> list[str]:
    items = [item for item in news if item.group == group][:limit]
    if not items:
        return ["- 暂无成功抓取的新闻。"]

    lines: list[str] = []
    for item in items:
        time_text = item.time or "时间待确认"
        lines.append(f"- 标题：{item.title}")
        lines.append(f"  - 时间：{time_text}")
        lines.append(f"  - 来源：{item.source}")
        lines.append(f"  - 链接：{item.link}")
        lines.append(f"  - 摘要：{item.summary}")
    return lines


def market_line(quote_obj: MarketQuote) -> str:
    if quote_obj.latest is None:
        return f"- {quote_obj.name}：未取到（{quote_obj.note or quote_obj.symbol}）"
    if quote_obj.name == "美债10Y":
        bps = quote_obj.change * 100 if quote_obj.change is not None else None
        bps_text = f"{bps:+.1f}bp" if bps is not None else "变动未取到"
        return f"- {quote_obj.name}：{format_number(quote_obj.latest, 3)}{quote_obj.unit}（{bps_text}，{quote_obj.symbol}）"
    return f"- {quote_obj.name}：{format_number(quote_obj.latest, 2)}{quote_obj.unit}（{format_pct(quote_obj.pct_change)}，{quote_obj.symbol}）"


def set_run_font(run, size: int | None = None, bold: bool = False, color: str | None = None) -> None:
    """统一设置中英文字体，保证 Word 中中文显示稳定。"""
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def setup_doc_styles(doc: Document) -> None:
    """应用 compact_reference_guide：紧凑，但保留正常阅读节奏。"""
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    list_style = doc.styles["List Bullet"]
    list_style.font.name = "Calibri"
    list_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    list_style.font.size = Pt(11)
    list_style.paragraph_format.left_indent = Inches(0.375)
    list_style.paragraph_format.first_line_indent = Inches(-0.188)
    list_style.paragraph_format.space_after = Pt(4)
    list_style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_run = header.add_run(f"{FRAMEWORK_VERSION} · A股盘前决策简报")
    set_run_font(header_run, size=8, color="6B7280")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{FRAMEWORK_VERSION} 盘前决策简报  ·  仅供纪律复核  ·  第 ")
    set_run_font(run, size=8, color="6B7280")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def set_cell_shading(cell, fill: str) -> None:
    """设置表格单元格底色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in: list[float]) -> None:
    """固定表格宽度，使 tblW、tblGrid 与单元格宽度一致。"""
    widths = [int(round(value * 1440)) for value in widths_in]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), "120")
    tbl_indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[index])
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    run.append(props)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def mark_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """写入表格单元格文本并统一格式。"""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    set_run_font(run, size=9, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_bullet(doc: Document, text: str) -> None:
    """添加一个普通项目符号段落。"""
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run, size=10)


def add_framework_snapshot(doc: Document, snapshot: dict[str, str], anchor_action: str, risk_preference: str) -> None:
    """Front-load the few facts that determine whether the reader may act."""
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    items = [
        ("市场权限", anchor_action),
        ("双锚", snapshot.get("双锚", "待确认")),
        ("情绪温度", snapshot.get("情绪", "待确认")),
        ("情绪可靠性", snapshot.get("情绪可靠性", "待确认")),
        ("隔夜风险", risk_preference),
        ("标准/半额首批", f"{snapshot.get('标准首批', '0')} / {snapshot.get('半额复核', '0')}"),
        ("质量待补", snapshot.get("质量待补", "待确认")),
        ("现金安全垫", snapshot.get("现金安全垫", "待确认")),
        ("第一年目标", snapshot.get("第一年目标占比", "待确认")),
        ("第一年当前", snapshot.get("第一年当前占比", "待确认")),
        ("第一年完成率", snapshot.get("第一年完成率", "待确认")),
        ("第一年资金缺口", snapshot.get("第一年资金缺口", "待确认")),
    ]
    for index, (label, value) in enumerate(items):
        cell = table.rows[index // 4].cells[index % 4]
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_run = paragraph.add_run(f"{label}\n")
        set_run_font(label_run, size=8, color="64748B")
        value_run = paragraph.add_run(str(value))
        set_run_font(value_run, size=11, bold=True, color="0B2545")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_shading(cell, "F3F7FB" if index < 4 else "EAF2F8")
    mark_table_header(table.rows[0])
    set_table_geometry(table, [1.625, 1.625, 1.625, 1.625])


def add_news_group(doc: Document, title: str, news: list[NewsItem], group: str, limit: int = 3) -> None:
    """写入一组新闻，保留标题、时间、来源、链接、摘要。"""
    doc.add_heading(title, level=2)
    items = [item for item in news if item.group == group][:limit]
    if not items:
        add_bullet(doc, "本组暂无已验证新闻，不据此生成方向判断。")
        return

    for item in items:
        heading = doc.add_paragraph()
        run = heading.add_run(item.title)
        set_run_font(run, size=10, bold=True, color="0B2545")

        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"时间：{item.time or '时间待确认'}    来源：{item.source}")
        set_run_font(meta_run, size=9, color="555555")

        if item.summary:
            summary = doc.add_paragraph()
            summary_run = summary.add_run(f"摘要：{item.summary}")
            set_run_font(summary_run, size=9)

        if item.link:
            link = doc.add_paragraph()
            label_run = link.add_run("来源：")
            set_run_font(label_run, size=8, color="555555")
            add_hyperlink(link, "打开原文", item.link)


def add_market_table(doc: Document, quotes: list[MarketQuote]) -> None:
    """写入隔夜行情表。"""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["指标", "最新", "变动", "数据时间"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    mark_table_header(table.rows[0])

    if not quotes:
        row = table.add_row().cells
        set_cell_text(row[0], "离线校验")
        set_cell_text(row[1], "待联网更新")
        set_cell_text(row[2], "不形成盘前方向判断")
        set_cell_text(row[3], "未联网")
    for quote_obj in quotes:
        row = table.add_row().cells
        if quote_obj.latest is None:
            latest_text = "未取到"
            change_text = quote_obj.note or quote_obj.symbol
        elif quote_obj.name == "美债10Y":
            bps = quote_obj.change * 100 if quote_obj.change is not None else None
            latest_text = f"{format_number(quote_obj.latest, 3)}{quote_obj.unit}"
            change_text = f"{bps:+.1f}bp（{quote_obj.symbol}）" if bps is not None else f"变动未取到（{quote_obj.symbol}）"
        else:
            latest_text = f"{format_number(quote_obj.latest, 2)}{quote_obj.unit}"
            change_text = f"{format_pct(quote_obj.pct_change)}（{quote_obj.symbol}）"

        set_cell_text(row[0], quote_obj.name)
        set_cell_text(row[1], latest_text)
        set_cell_text(row[2], change_text)
        set_cell_text(row[3], quote_obj.data_time or "未取到")
    set_table_geometry(table, [1.05, 1.45, 2.10, 1.90])


def add_position_table(doc: Document, position_impacts: list[dict[str, str]]) -> None:
    """写入持仓影响表。"""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["持仓", "新闻/行情影响", "框架判断", "动作"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
    mark_table_header(table.rows[0])

    rows = position_impacts or [
        {
            "持仓": "未读取到持仓",
            "新闻/行情影响": "请先生成最新 V2.8.5 每日行情输出",
            "框架判断": "暂停",
            "动作": "暂停",
        }
    ]
    for item in rows:
        row = table.add_row().cells
        set_cell_text(row[0], item["持仓"])
        set_cell_text(row[1], item["新闻/行情影响"])
        set_cell_text(row[2], item["框架判断"])
        set_cell_text(row[3], item["动作"])
    set_table_geometry(table, [1.25, 2.35, 2.10, 0.80])


def add_execution_plan_table(doc: Document, execution_plan: pd.DataFrame) -> None:
    """展示当日动作预算，不把复核资格写成交易指令。"""
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["优先级", "动作类型", "标的", "状态/依据"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "E8EEF5")
    mark_table_header(table.rows[0])
    if execution_plan.empty:
        items = [{"优先级": 1, "动作类型": "观察", "标的": "无", "状态": "质量、情绪、买点或仓位证据不足"}]
    else:
        items = []
        for _, row in execution_plan.head(6).iterrows():
            annual_gap = to_float(row.get("年度资金缺口"))
            annual_completion = to_float(row.get("年度完成率"))
            annual_note = ""
            if annual_gap is not None:
                annual_note = f"；年度缺口{annual_gap:,.0f}元"
            if annual_completion is not None:
                annual_note += f"，完成{annual_completion:.1%}"
            items.append(
                {
                    "优先级": row.get("优先级", ""),
                    "动作类型": row.get("动作类型", ""),
                    "标的": row.get("标的", ""),
                    "状态": f"{row.get('状态', '')}；{clean_text(row.get('依据', ''), 120)}{annual_note}",
                }
            )
    for item in items:
        row = table.add_row().cells
        set_cell_text(row[0], item["优先级"])
        set_cell_text(row[1], item["动作类型"])
        set_cell_text(row[2], item["标的"])
        set_cell_text(row[3], item["状态"])
    set_table_geometry(table, [0.70, 1.30, 1.75, 2.75])


def add_first_year_table(doc: Document, first_year: pd.DataFrame) -> None:
    """展示第一年全资产配置进度；缺口只作规划，不构成交易指令。"""
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["配置项", "目标占比", "动态目标", "当前金额", "资金缺口", "进度"]
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "E8EEF5")
    mark_table_header(table.rows[0])
    if first_year.empty:
        rows = [["未读取到配置表", "—", "—", "—", "—", "待补"]]
    else:
        rows = []
        for _, item in first_year.iterrows():
            rows.append(
                [
                    clean_text(item.get("配置项", ""), 28),
                    f"{to_float(item.get('年度目标占比')) or 0:.1%}",
                    f"{to_float(item.get('按当前全资产目标金额')) or 0:,.0f}",
                    f"{to_float(item.get('最新持仓金额')) or 0:,.0f}",
                    f"{to_float(item.get('年度资金缺口')) or 0:,.0f}",
                    str(item.get("进度状态", "")),
                ]
            )
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)
    set_table_geometry(table, [1.70, 0.70, 1.05, 1.00, 1.05, 1.00])


def build_data_audit_lines(
    sheets: dict[str, pd.DataFrame],
    quotes: list[MarketQuote],
    news: list[NewsItem],
    source_notes: list[str],
) -> list[str]:
    """Build a factual completeness trail for every report data class."""
    lines: list[str] = []
    dashboard_path = find_latest_dashboard_xlsx()
    dashboard = get_sheet_with_fallback(sheets, *PORTFOLIO_SUMMARY_SHEET_NAMES)
    dashboard_map: dict[str, Any] = {}
    if not dashboard.empty and {"项目", "内容"}.issubset(dashboard.columns):
        dashboard_map = dict(zip(dashboard["项目"].astype(str), dashboard["内容"]))
    if dashboard_path:
        mtime = datetime.fromtimestamp(dashboard_path.stat().st_mtime, SHANGHAI_TZ).strftime("%Y-%m-%d %H:%M:%S")
        lines.append(
            f"本地仪表盘：{dashboard_path.name}；文件修改时间 {mtime}；"
            f"内部生成时间 {dashboard_map.get('生成时间', '未取到')}；"
            f"行情数据时间 {dashboard_map.get('行情数据时间', '未取到')}"
        )
    else:
        lines.append("本地仪表盘：未找到每日行情输出文件")

    available_quotes = [quote_obj for quote_obj in quotes if quote_obj.latest is not None]
    missing_quotes = [quote_obj.name for quote_obj in quotes if quote_obj.latest is None]
    quote_times = "；".join(
        f"{quote_obj.name} {quote_obj.data_time or '时间未取到'}"
        for quote_obj in available_quotes
    )
    lines.append(
        f"隔夜/全球行情：{len(available_quotes)}/{len(quotes)}项有值"
        f"{'；缺失 ' + '、'.join(missing_quotes) if missing_quotes else ''}；{quote_times or '无数据时间'}"
    )
    for quote_obj in quotes:
        if quote_obj.note:
            lines.append(f"行情源追踪：{quote_obj.name}：{quote_obj.note}")

    group_counts: dict[str, int] = {}
    group_sources: dict[str, set[str]] = {}
    for item in news:
        group_counts[item.group] = group_counts.get(item.group, 0) + 1
        group_sources.setdefault(item.group, set()).add(item.source)
    parsed_news_times = [parse_datetime_or_none(item.time) for item in news if item.time]
    parsed_news_times = [value for value in parsed_news_times if value is not None]
    latest_news_time = max(parsed_news_times).strftime("%Y-%m-%d %H:%M") if parsed_news_times else "未取到"
    missing_news_time = sum(1 for item in news if not item.time)
    lines.append(
        "新闻完整性："
        + "；".join(
            f"{group} {group_counts.get(group, 0)}条/"
            f"{'、'.join(sorted(group_sources.get(group, set()))) or '无成功来源'}"
            for group in ["官方源", "国际源", "国内快讯"]
        )
        + f"；最新显示时间 {latest_news_time}；时间缺失 {missing_news_time}条"
    )
    lines.append(
        f"宏观/政策：官方源 {group_counts.get('官方源', 0)}条，"
        f"来源 {('、'.join(sorted(group_sources.get('官方源', set()))) or '无')}"
    )

    long_watch = sheets.get("长期跟踪个股", pd.DataFrame())
    if long_watch.empty:
        lines.append("长期跟踪个股：未读取到数据，本节不得用于盘前判断")
    else:
        tracked_total = len(long_watch)
        latest_available = int(long_watch["Latest"].notna().sum()) if "Latest" in long_watch.columns else 0
        pct_available = int(long_watch["PctChg"].notna().sum()) if "PctChg" in long_watch.columns else 0
        quote_time_available = int(long_watch["QuoteTime"].notna().sum()) if "QuoteTime" in long_watch.columns else 0
        quality_available = int(long_watch["质量评分"].notna().sum()) if "质量评分" in long_watch.columns else 0
        gate_available = int(long_watch["买点灯号"].notna().sum()) if "买点灯号" in long_watch.columns else 0
        lines.append(
            f"长期跟踪个股：{tracked_total}只；Latest {latest_available}/{tracked_total}；"
            f"PctChg {pct_available}/{tracked_total}；QuoteTime {quote_time_available}/{tracked_total}；"
            f"质量评分 {quality_available}/{tracked_total}；买点灯号 {gate_available}/{tracked_total}"
        )

    investment_profile = sheets.get("Investment_Profile", pd.DataFrame())
    data_sources = sheets.get("Data_Sources", pd.DataFrame())
    profile_rows = len(investment_profile) if not investment_profile.empty else 0
    source_rows = len(data_sources) if not data_sources.empty else 0
    lines.append(f"配置与来源：Investment_Profile {profile_rows}行；Data_Sources {source_rows}行")

    core_sheet_groups = [
        ("今日决策", ("01_今日决策", "Decision_Center")),
        ("组合总览", PORTFOLIO_SUMMARY_SHEET_NAMES),
        ("行情数据", ("Market_Data",)),
        ("双锚", ("Double_Anchor",)),
        ("情绪", ("Emotion",)),
        ("质量评分", ("Quality_Score",)),
        ("组合暴露", ("Exposure",)),
        ("买入候选", ("03_买入候选", "Buy_Filter")),
        ("持仓", ("Positions",)),
        ("持仓风险", POSITION_RISK_SHEET_NAMES),
        ("年度配置", FIRST_YEAR_SHEET_NAMES),
        ("今日动作", ("02_今日动作", "Execution_Plan")),
        ("校验", ("Checks",)),
        ("规则", ("Framework_Rules",)),
        ("长期跟踪个股", ("长期跟踪个股",)),
        ("投资画像", ("Investment_Profile",)),
        ("数据来源", ("Data_Sources",)),
    ]
    empty_sheets = [label for label, names in core_sheet_groups if get_sheet_with_fallback(sheets, *names).empty]
    lines.append(
        f"其他决策数据：{len(core_sheet_groups) - len(empty_sheets)}/{len(core_sheet_groups)}个核心工作表非空"
        f"{'；空表 ' + '、'.join(empty_sheets) if empty_sheets else ''}"
    )
    lines.extend(f"来源追踪：{note}" for note in source_notes)
    if not source_notes:
        lines.append("来源追踪：本次所有配置源首次抓取均返回近3日有效数据")
    return lines


def write_report(
    quotes: list[MarketQuote],
    risk_preference: str,
    risk_notes: list[str],
    news: list[NewsItem],
    news_warnings: list[str],
    position_impacts: list[dict[str, str]],
    theme_view: dict[str, str],
    anchor_light: str,
    anchor_action: str,
    buy_filter_risks: list[str],
    discipline: dict[str, str],
    snapshot: dict[str, str],
    execution_plan: pd.DataFrame,
    checks: pd.DataFrame,
    data_audit_lines: list[str],
    first_year: pd.DataFrame,
) -> Path:
    """写入 Word 报告。"""
    now = datetime.now()
    output_path = OUTPUT_DIR / f"{FRAMEWORK_VERSION}_A股盘前决策简报_{now.strftime('%Y%m%d_%H%M%S')}.docx"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    setup_doc_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("A股盘前决策简报")
    set_run_font(title_run, size=22, bold=True, color="0B2545")

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"{FRAMEWORK_VERSION} · 生成时间：{now.strftime('%Y-%m-%d %H:%M:%S')}")
    set_run_font(meta_run, size=10, color="555555")

    boundary = doc.add_paragraph()
    boundary.paragraph_format.space_before = Pt(4)
    boundary.paragraph_format.space_after = Pt(10)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FDECEC")
    boundary._p.get_or_add_pPr().append(shading)
    boundary_run = boundary.add_run("纪律边界：市场权限 ≠ 执行信号；质量、情绪、买点、仓位任一未通过即不执行。本报告不连接券商、不自动交易。")
    set_run_font(boundary_run, size=10, bold=True, color="9B1C1C")

    doc.add_heading("1. 决策摘要", level=1)
    conclusion = doc.add_paragraph()
    buy_ready = int(to_float(snapshot.get("标准首批")) or 0) + int(to_float(snapshot.get("半额复核")) or 0)
    portfolio_gate = "暂停新增，仅处理复审与观察" if buy_ready == 0 else "存在条件式候选，仍须盘中逐项确认"
    conclusion_run = conclusion.add_run(
        f"结论：市场环境权限为“{anchor_action}”，组合层面“{portfolio_gate}”，隔夜风险偏好“{risk_preference}”；"
        f"质量待补{snapshot.get('质量待补', '待确认')}只，校验状态“{snapshot.get('校验状态', '待确认')}”。"
    )
    set_run_font(conclusion_run, size=11, bold=True, color="0B2545")
    add_framework_snapshot(doc, snapshot, anchor_action, risk_preference)
    add_bullet(doc, "执行优先级：风控 > 清旧 > 补新 > 再平衡；单日最多一个可检查方向和两个复审方向。")

    doc.add_heading("2. 今日动作预算", level=1)
    add_execution_plan_table(doc, execution_plan)

    doc.add_heading("3. 第一年度配置进度", level=1)
    add_first_year_table(doc, first_year)
    add_bullet(doc, "年度资金缺口只表示长期配置上限；不得突破质量、双锚、情绪、买点或证券账户集中度约束。")

    doc.add_heading("4. 隔夜行情与风险偏好", level=1)
    add_market_table(doc, quotes)
    add_bullet(doc, f"风险偏好：{risk_preference}")
    if risk_notes:
        add_bullet(doc, f"判断依据：{'；'.join(risk_notes)}")

    doc.add_heading("5. 当前持仓风险与动作", level=1)
    add_position_table(doc, position_impacts)

    doc.add_heading("6. 主线、质量与买点门槛", level=1)
    add_bullet(doc, f"半导体设备：{theme_view.get('半导体', '观察')}")
    add_bullet(doc, f"AI算力：{theme_view.get('AI算力', '观察')}")
    add_bullet(doc, f"机器人：{theme_view.get('机器人', '观察')}")
    add_bullet(doc, f"创新药：{theme_view.get('创新药', '观察')}")
    add_bullet(doc, f"双锚综合：{anchor_light}；盘前动作门槛：{anchor_action}")
    for item in buy_filter_risks:
        add_bullet(doc, item)

    doc.add_heading("7. 重大新闻与政策线索", level=1)
    if news:
        add_news_group(doc, "官方源", news, "官方源")
        add_news_group(doc, "国际源", news, "国际源")
        add_news_group(doc, "国内快讯", news, "国内快讯")
    else:
        add_bullet(doc, "未取得已验证新闻；本节不提供替代性猜测，盘前判断仅基于行情仪表盘。")

    doc.add_heading("8. 今日执行纪律", level=1)
    add_bullet(doc, f"禁止追高：{discipline['禁止追高']}")
    add_bullet(doc, f"禁止自动挂单：{discipline['禁止自动挂单']}")
    add_bullet(doc, f"需要复审：{discipline['需要复审']}")
    add_bullet(doc, f"可观察：{discipline['可观察']}")

    doc.add_heading("9. 数据完整性与校验", level=1)
    add_bullet(doc, snapshot.get("数据状态", "条件式"))
    for line in data_audit_lines:
        add_bullet(doc, line)
    if not checks.empty and "状态" in checks.columns:
        for _, row in checks[checks["状态"].astype(str) != "OK"].head(6).iterrows():
            add_bullet(doc, f"{row.get('检查项', '校验')}：{row.get('状态', '检查')}；{row.get('修复建议', '')}")
    if news_warnings and any("离线" in warning for warning in news_warnings):
        add_bullet(doc, "离线校验模式：隔夜行情与新闻未联网更新，正式盘前不得沿用本节数据。")
    add_bullet(doc, "ETF份额/IOPV、个股财务/估值未接入时，只能条件式复核，不得自动判定为可执行。")

    doc.save(output_path)
    return output_path


def main() -> int:
    print("=" * 60)
    print(f"{FRAMEWORK_VERSION} A股盘前决策简报")
    print("=" * 60)

    offline = "--offline" in sys.argv

    print("正在读取 V2.8.5 每日行情输出 ...")
    sheets = read_dashboard_workbook()

    if offline:
        print("离线校验模式：跳过隔夜行情与新闻抓取。")
        quotes = []
        risk_preference, risk_notes = "中性", ["离线校验模式，隔夜风险待联网更新"]
        news, news_warnings = [], ["离线校验模式：未抓取新闻"]
    else:
        if requests is None or date_parser is None:
            raise RuntimeError("缺少联网依赖，请先安装 premarket_report/requirements.txt，或使用 --offline 校验文档。")
        print("正在抓取隔夜行情 ...")
        quotes = fetch_market_quotes()
        risk_preference, risk_notes = judge_risk_preference(quotes)

        print("正在抓取新闻源 ...")
        news, news_warnings = fetch_news()

    anchor_light, anchor_action = get_anchor_gate(sheets, risk_preference)
    buy_filter_risks = get_buy_filter_risks(sheets)
    position_impacts = build_position_impacts(sheets, news, risk_preference, anchor_action)
    theme_view = build_theme_view(news, risk_preference, anchor_action)
    discipline = build_today_discipline(position_impacts, buy_filter_risks, risk_preference, anchor_action, theme_view)
    snapshot = build_framework_snapshot(sheets)
    execution_plan = get_sheet_with_fallback(sheets, "02_今日动作", "Execution_Plan")
    checks = sheets.get("Checks", pd.DataFrame())
    first_year = get_sheet_with_fallback(sheets, *FIRST_YEAR_SHEET_NAMES)
    data_audit_lines = build_data_audit_lines(sheets, quotes, news, news_warnings)

    output_path = write_report(
        quotes=quotes,
        risk_preference=risk_preference,
        risk_notes=risk_notes,
        news=news,
        news_warnings=news_warnings,
        position_impacts=position_impacts,
        theme_view=theme_view,
        anchor_light=anchor_light,
        anchor_action=anchor_action,
        buy_filter_risks=buy_filter_risks,
        discipline=discipline,
        snapshot=snapshot,
        execution_plan=execution_plan,
        checks=checks,
        data_audit_lines=data_audit_lines,
        first_year=first_year,
    )

    print()
    print(f"已生成盘前报告：{output_path}")
    if offline:
        print("提示：这是离线校验文档；正式盘前使用请运行 python main.py 联网更新。")
    elif news_warnings:
        print("注意：存在新闻源重试/备用源记录，报告末尾已逐项列出。")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except KeyboardInterrupt:
        print("\n用户中断。")
        raise SystemExit(1)
    except Exception as exc:
        print(f"\n【程序出错】{exc}")
        raise SystemExit(1)
