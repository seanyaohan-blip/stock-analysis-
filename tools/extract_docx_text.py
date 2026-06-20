from __future__ import annotations

import sys
from pathlib import Path

from docx import Document


def extract(path: Path) -> None:
    doc = Document(path)
    print(f"### {path.name}")
    print("PARAGRAPHS")
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            print(f"[{paragraph.style.name}] {text}")
    print(f"TABLES {len(doc.tables)}")
    for index, table in enumerate(doc.tables, start=1):
        print(f"[TABLE {index}]")
        for row in table.rows:
            print(" | ".join(cell.text.strip().replace("\n", " / ") for cell in row.cells))


def main() -> int:
    for value in sys.argv[1:]:
        extract(Path(value))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
