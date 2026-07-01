from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "投资决策分析框架 V2.8.5.docx"

DECISION_CHAIN = (
    "产业主线 → 主线周期 → 双锚系统 → 资产质量 → 量化验证层 → "
    "爆量/回踩/量质吸结构 → 买点过滤器 → 情绪温度 → 认知防错层 → 仓位执行 → 复盘迭代。"
)


def set_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10.5)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9)
    doc.add_paragraph()


def already_has_quant_section(doc: Document) -> bool:
    return any("V2.8.5-Q 量化验证与执行纪律模块" in paragraph.text for paragraph in doc.paragraphs)


def already_has_dividend_section(doc: Document) -> bool:
    return any("红利现金流层与红利个股买点规则" in paragraph.text for paragraph in doc.paragraphs)


def update_existing_text(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "主线周期、全资产目标与趋势买点增强合并版":
            set_text(paragraph, "主线周期、量化验证、全资产目标与趋势买点增强合并版")
        elif text.startswith("合并更新日期："):
            set_text(paragraph, "合并更新日期：2026年6月25日")
        elif text == "V2.8.5 新增三大模块：":
            set_text(paragraph, "V2.8.5 当前增强为五大模块：")
        elif text == "V2.8.5 当前增强为五大模块：":
            set_text(paragraph, "V2.8.5 当前增强为六大模块：")
        elif "趋势买点增强模块" == text:
            continue
        elif "比亚迪 / 绿电 / 储能 / 中望软件" in text:
            set_text(paragraph, text.replace("比亚迪 / 绿电 / 储能 / 中望软件", "比亚迪 / 绿电 / 中望软件"))
        elif "比亚迪、绿电、储能" in text:
            set_text(paragraph, text.replace("比亚迪、绿电、储能", "比亚迪、绿电"))
        elif "云计算/绿电/比亚迪/中望" in text:
            set_text(paragraph, text.replace("云计算/绿电/比亚迪/中望", "云计算/绿电/比亚迪/中望"))
        elif text == "产业主线 → 主线周期 → 双锚系统 → 资产质量 → 爆量/回踩/量质吸结构 → 买点过滤器 → 情绪温度 → 认知防错层 → 仓位执行 → 次日验证。":
            set_text(paragraph, DECISION_CHAIN)
        elif text.startswith("产业主线") and "主线周期" in text and "仓位纪律" in text:
            set_text(paragraph, DECISION_CHAIN)
        elif text.startswith("最终总纲："):
            set_text(
                paragraph,
                "最终总纲：主线决定方向，双锚决定权限，资产质量决定标的，量化验证决定信号质量，趋势买点决定时机，买点过滤器决定能否执行，情绪温度决定是否追，认知防错层决定是否被偏见误导，仓位系统决定买多少，复盘迭代决定框架是否持续进化。",
            )
        elif text == "第二十部分：V2.8.5-M 多元思维模型认知防错层":
            set_text(paragraph, "第二十部分：V2.8.5-QM 多元思维模型认知防错层")


def append_quant_section(doc: Document) -> None:
    doc.add_heading("第二十一部分：V2.8.5-Q 量化验证与执行纪律模块", level=1)

    doc.add_heading("一、模块定位", level=2)
    doc.add_paragraph(
        "V2.8.5-Q 不是把框架改成纯量化系统，而是在原有产业主线、双锚、资产质量、趋势买点和仓位纪律之上，增加数据验证层、执行纪律层和风险监控层。"
    )
    doc.add_paragraph(
        "最核心的原则是：主观框架定方向，量化模块验信号；主观判断提出假设，数据负责验证假设，规则负责约束执行，复盘负责修正模型。"
    )
    doc.add_paragraph(
        "量化模块不得单独产生买入信号；资金流入、单日放量、研报催化或短期上涨，都必须经过价格结构、资产角色、正期望、情绪温度和仓位纪律复核。"
    )

    doc.add_heading("二、总决策链条更新", level=2)
    doc.add_paragraph(DECISION_CHAIN)

    doc.add_heading("三、正期望与量化打分", level=2)
    doc.add_paragraph("量化验证不追求固定的高胜率低赔率，而是追求可解释、可复盘、经过条件检验的正期望。")
    doc.add_paragraph("期望收益 = 胜率 × 平均盈利 - 失败率 × 平均亏损。")

    add_table(
        doc,
        ["指标", "用途"],
        [
            ["上证/沪深300/中证A500涨跌幅", "总量锚，判断系统权限"],
            ["创业板指/科创50涨跌幅", "成长锚，判断成长风格权限"],
            ["两市成交额与成交额变化率", "流动性和放量/缩量判断"],
            ["上涨家数占比", "市场广度，避免少数权重掩盖真实弱势"],
            ["ETF份额/资金方向", "验证资金是否认可主线"],
            ["折溢价/流动性", "避免因交易结构问题追高或误判"],
        ],
    )

    add_table(
        doc,
        ["ETF买点指标", "分值", "说明"],
        [
            ["分时结构", "0-2", "回踩不破、收盘不接近日内低点优先"],
            ["量价关系", "0-2", "放量上涨、缩量企稳优于放量下跌或滞涨"],
            ["ETF份额/筹码", "0-2", "稳定或净流入加分，连续赎回扣分"],
            ["折溢价/估值", "0-1", "折价或轻微溢价优于高溢价追买"],
            ["龙头同步", "0-2", "核心龙头共振强于后排乱涨"],
            ["次日验证", "0-1", "次日不大幅回吐或重新转强"],
        ],
    )

    add_table(
        doc,
        ["量化验证分", "动作"],
        [
            ["8-10分", "标准买点，可进入人工复核"],
            ["6-7分", "黄灯观察，只能小额或半额复核"],
            ["4-5分", "不买，等待条件改善"],
            ["低于4分", "风险复审，优先检查是否退潮或数据失真"],
        ],
    )

    doc.add_heading("四、资产角色差异化模板", level=2)
    add_table(
        doc,
        ["资产角色", "适用对象", "买点规则", "风控边界"],
        [
            ["核心成长ETF", "159516、159381等", "只买缩量回踩不破+龙头同步+次日转强", "连续上涨不追；净赎回、放量滞涨或龙头破位暂停新增"],
            ["低波权益ETF", "159338、512890、510210", "按配置节奏分批，不在指数大涨日追买", "系统性破位、权重股集体破位或重复暴露过高时复核"],
            ["弹性低波ETF", "511180", "每次小额分批，缩量回调或横盘企稳优先", "转债估值过热、利率急升或权益风险扩散时暂停新增"],
            ["纯防御ETF", "511010、现金/货基", "只做安全垫和弹药，不追单日收益", "权益仓提高时必须保留，不因短期跑输科技而切走"],
            ["降级/遗留仓", "云计算、绿电、比亚迪、中望", "不使用买点新增，只看反弹压缩或退出复审", "沉没成本不构成补仓理由，继续弱于核心资产则降级"],
            ["卫星验证个股", "AI/半导体核心个股篮子", "质量≥8、主线未退潮、回踩不破、首仓小", "先定义证伪和最大损失，跌破关键位快速复审"],
        ],
    )

    doc.add_heading("五、当前组合落地", level=2)
    add_table(
        doc,
        ["标的/类别", "量化模块用法", "当前纪律"],
        [
            ["159516 半导体设备ETF", "防止看对方向但买错位置；等待回踩确认提高胜率", "第一核心候选，连续上涨不追"],
            ["159381 创业板人工智能ETF", "防止被AI叙事和研报催化带偏", "AI硬件增强仓，优先级低于159516"],
            ["159338 中证A500ETF", "作为低波权益底仓，不作为追涨工具", "守住关键位、缩量、权重股止跌后分批"],
            ["562500 机器人ETF", "防止盈利仓冲动加仓", "持有不加，等止跌三部曲"],
            ["红利低波/可转债/国债", "降低组合波动，提供弹性缓冲和安全垫", "按角色分批，不与科技仓比单日弹性"],
            ["云计算/绿电/比亚迪/中望", "用机会成本和沉没成本模型约束补仓冲动", "不补仓，反弹压缩或退出复审"],
            ["AI/半导体个股篮子", "只做产业链弹性验证", "首仓0.3%-0.5%，未补完整评分前只观察"],
        ],
    )

    doc.add_heading("六、最终原则", level=2)
    doc.add_paragraph(
        "V2.8.5-Q 吸收量化投资的规则化、概率化和纪律化，但不变成纯量化系统；核心仍是产业主线、双锚、资产质量、买点过滤和全资产仓位管理。"
    )
    doc.add_paragraph(
        "一句话：主观判断方向，量化验证买点，纪律控制仓位，复盘修正模型。"
    )


def append_dividend_section(doc: Document) -> None:
    doc.add_heading("第二十二部分：红利现金流层与红利个股买点规则", level=1)

    doc.add_heading("一、模块定位", level=2)
    doc.add_paragraph(
        "红利个股在 V2.8.5-QM 中归入低波现金流层，不属于主线进攻仓。它的任务是降低组合波动、提供现金流、在科技主线拥挤时做风格平衡，而不是替代半导体设备、AI硬件和机器人等核心成长收益发动机。"
    )
    doc.add_paragraph(
        "红利配置的重点不是寻找最高股息率，而是排除高息陷阱后，选择长期分红稳定、经营现金流能覆盖分红、估值具备安全垫且价格处在缩量回踩或横盘止跌结构中的公司。"
    )

    doc.add_heading("二、资产角色与年度配置", level=2)
    add_table(
        doc,
        ["资产/组合", "框架定位", "第一阶段目标", "执行含义"],
        [
            ["红利低波ETF", "低波权益ETF底仓", "全资产约5%", "用ETF承接红利风格底仓，按配置节奏分批，不当现金安全垫"],
            ["长期稳定分红个股篮子", "低波现金流个股", "全资产约5%", "先建立观察池，满足红利买点后慢建，不追高"],
            ["红利个股+红利低波合计", "组合稳定器", "初始3%–5%，第一阶段6%–8%", "稳定后可提高到10%–15%，但不得挤占核心成长仓和现金安全垫"],
        ],
    )

    doc.add_heading("三、行业和标的观察池", level=2)
    add_table(
        doc,
        ["优先级", "方向", "候选", "框架判断"],
        [
            ["第一优先", "公用事业/运营商/水电", "长江电力、中国移动", "现金流稳定性更强，适合作为红利底仓观察"],
            ["第二优先", "国有大行", "工商银行、建设银行", "分红规模和低波属性较强，但需跟踪净息差与资产质量"],
            ["第三优先", "能源红利", "中国神华", "股息吸引力强，但周期属性更高，只做红利增强"],
            ["第四优先", "质量消费红利", "美的集团", "分红与经营质量兼具，需复核估值和增长质量"],
        ],
    )

    doc.add_heading("四、红利买入绿灯", level=2)
    add_table(
        doc,
        ["检查项", "绿灯标准"],
        [
            ["分红连续性", "至少5年稳定分红，最好10年以上"],
            ["分红率", "30%–75%较健康，周期股高于80%需谨慎"],
            ["现金流", "经营现金流能覆盖分红，不依赖举债分红"],
            ["盈利趋势", "净利润不应连续明显下滑"],
            ["股息率安全垫", "银行4.5%+，能源5%+，公用事业3.5%–4%+，消费3%–4%+更有吸引力"],
            ["估值", "PE/PB处于历史中低位"],
            ["技术形态", "缩量回调、横盘止跌、未放量破位"],
            ["行业风险", "没有明显利润下修、政策冲击或资本开支恶化"],
        ],
    )

    doc.add_heading("五、高息陷阱一票否决", level=2)
    add_table(
        doc,
        ["否决项", "处理"],
        [
            ["分红率超过100%，且盈利或现金流变差", "不买"],
            ["股息率很高只是因为股价暴跌", "不买"],
            ["高分红同时大额举债或资本开支明显增加", "不买"],
            ["银行不良率上行、净息差恶化超预期", "不买或降级"],
            ["煤炭/石化处在盈利高点，市场用高股息诱导追买", "只观察，不追"],
            ["公用事业出现电价、收费政策、资本开支重大变化", "复审"],
            ["买入理由只是“快分红了”", "不买"],
        ],
    )

    doc.add_heading("六、建仓与复审纪律", level=2)
    add_table(
        doc,
        ["项目", "规则"],
        [
            ["组合初始仓位", "红利个股+红利低波合计先建全资产3%–5%"],
            ["第一阶段目标", "条件成熟后提高到6%–8%"],
            ["稳定后上限", "10%–15%，不得超过低波权益层总额度"],
            ["A档单只", "初始0.5%–1%，成熟1.5%–3%"],
            ["B档单只", "初始0.3%–0.8%，成熟1%–2%"],
            ["跌10%", "复核是否只是风格下跌"],
            ["跌15%", "强制复审分红逻辑"],
            ["跌20%且利润/现金流恶化", "至少减半"],
            ["公司削减分红或现金流连续低于分红", "降级或退出"],
        ],
    )


def main() -> None:
    doc = Document(DOCX_PATH)
    update_existing_text(doc)
    if not already_has_quant_section(doc):
        append_quant_section(doc)
    if not already_has_dividend_section(doc):
        append_dividend_section(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
